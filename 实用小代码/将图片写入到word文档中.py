from docx import Document

doc = Document()                #以默认模板建立文档对象
doc.add_paragraph("hello")
doc.add_picture('C:/Users/Administrator/Desktop/python_work/微信图片_20210820101807.png')
doc.save('a.docx')     # 读取a.docx文档，建立文档对象

